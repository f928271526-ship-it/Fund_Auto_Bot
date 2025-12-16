from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. 创建文档对象
doc = Document()

# --- 辅助函数：添加带格式的段落 ---
def add_para(text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    runner = p.add_run(text)
    runner.bold = bold
    runner.italic = italic
    runner.font.size = Pt(size)
    # 设置中文字体需要特殊处理，这里用默认字体，Word会自动匹配宋体/等线
    return p

def add_math_block(text):
    """添加类似数学公式的独立段落"""
    p = doc.add_paragraph()
    runner = p.add_run(text)
    runner.font.name = 'Cambria Math' # 尝试使用数学字体
    runner.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中显示

# ================= 报告内容开始 =================

# 1. 封面/标题信息
title = doc.add_heading('《人工智能导论》课程学习报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para('\n') # 空行
info_lines = [
    "—— 基于混合AI模型的数字金融量化决策研究",
    "",
    "课程名称：人工智能导论",
    "任课教师：[请填写老师姓名]",
    "学生姓名：符清华",
    "学号：52515056028",
    "专业：数字经济",
    "年级：2025级硕士研究生",
    "教材依据：《人工智能导论（第五版）》王万良"
]
for line in info_lines:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break() # 换页

# 2. 摘要
doc.add_heading('摘要', level=1)
abstract_text = (
    "数字经济时代，数据已成为核心生产要素。在金融市场中，交易数据呈现出高噪声、非线性及模糊性的特征，传统线性统计模型难以应对。本报告基于《人工智能导论》课程核心知识点，选取“模糊推理（Fuzzy Reasoning）”、“计算智能中的遗传算法（Genetic Algorithm）”以及“人工神经网络中的感知机（Perceptron）”三个模块进行深入总结。报告结合本人在数字经济专业研究中开发的“Fund_Auto_Bot”量化交易系统，详细阐述了上述算法在基金仓位控制、资产配置及趋势预测中的实际计算过程与应用分析，旨在探索 AI 技术赋能数字金融决策的有效路径。"
)
add_para(abstract_text)

# 3. 第一章
doc.add_heading('一、 核心知识点理论阐述', level=1)

doc.add_heading('1.1 模糊推理 (Fuzzy Reasoning)', level=2)
add_para("现实世界中，许多概念（如“市场过热”、“价格偏低”）具有模糊性，无法用二值逻辑（0或1）精确描述。模糊推理基于模糊集合论，通过隶属度函数将精确数值转化为模糊变量。")
add_para("推理流程：通常包含模糊化（Fuzzification）、模糊关系矩阵建立（利用直积法或取小运算）、合成推理（Max-Min Composition）以及解模糊（Defuzzification）四个步骤。")

doc.add_heading('1.2 计算智能：遗传算法的选择机制 (Selection in GA)', level=2)
add_para("遗传算法是模拟生物进化论的随机搜索算法。在种群进化的过程中，选择算子（Selection Operator）决定了哪些优良个体能遗传到下一代。")
add_para("轮盘赌选择法：这是最经典的选择策略。通过计算个体的适应度（Fitness）在总适应度中的占比（选择概率），构建累积概率分布。随机产生一个 [0,1] 的数，落入哪个区间即选择哪个个体，体现了“优胜劣汰”的概率机制。")

doc.add_heading('1.3 人工神经网络：感知机模型 (Perceptron)', level=2)
add_para("感知机是神经网络的基础，模拟神经元对输入信号的加权求和与激活过程。")
add_math_block("模型结构：y = f(∑ ωi * xi - θ)")
add_para("学习规则：这是一种有监督学习。通过比较模型输出与期望输出的误差，利用权重更新规则不断迭代修正，直至收敛。")
add_math_block("Δω = η * (y_expect - y_actual) * x")

# 4. 第二章 (核心计算部分)
doc.add_heading('二、 专业应用案例分析：AI 驱动的量化交易系统', level=1)
add_para("本部分结合 Python 量化实战数据（国泰证券、富国煤炭等标的），详细展示上述三种算法的计算细节。")

# 2.1 模糊推理
doc.add_heading('2.1 应用一：基于模糊矩阵合成的 RSI 动态仓位决策', level=2)
add_para("在量化交易中，RSI 指标与买入仓位之间的关系是模糊的。为了确定今日具体的买入比例，我采用了模糊关系矩阵合成的方法进行推理。")

add_para("【1. 定义论域与模糊向量】", bold=True)
add_para("设输入论域 U（RSI 状态）= {低, 中, 高}。")
add_para("设输出论域 V（仓位档位）= {轻仓(20%), 中仓(50%), 重仓(80%)}。这两组数值将作为解模糊时的加权因子 yj。")

add_para("【2. 第一步：建立模糊关系矩阵 R】", bold=True)
add_para("根据历史经验定义规则：“如果 RSI 低 (A)，则 仓位重 (B)”。")
add_para("设模糊向量 A（RSI 低）= [1.0, 0.4, 0.0] （表示非常确定的低位）。")
add_para("设模糊向量 B（仓位重）= [0.2, 0.5, 1.0] （表示倾向于重仓）。")
add_para("利用笛卡尔积（取小运算 min）计算模糊关系矩阵 R = A^T × B：")

# 模拟矩阵格式
add_math_block("R = [min(1.0,0.2)  min(1.0,0.5)  min(1.0,1.0)]")
add_math_block("    [min(0.4,0.2)  min(0.4,0.5)  min(0.4,1.0)]")
add_math_block("    [min(0.0,0.2)  min(0.0,0.5)  min(0.0,1.0)]")
add_math_block("= [[0.2, 0.5, 1.0], [0.2, 0.4, 0.4], [0.0, 0.0, 0.0]]")

add_para("【3. 第二步：模糊推理计算 B'】", bold=True)
add_para("假设今日国泰证券的 RSI 为 31.25，计算出今日的特征向量 A' = [0.8, 0.3, 0.1]。")
add_para("利用最大-最小合成法（Max-Min Composition）计算输出模糊向量 B' = A' ○ R：")
add_para("计算过程：")
add_math_block("b'1 = max(min(0.8,0.2), min(0.3,0.2), min(0.1,0.0)) = 0.2")
add_math_block("b'2 = max(min(0.8,0.5), min(0.3,0.4), min(0.1,0.0)) = 0.5")
add_math_block("b'3 = max(min(0.8,1.0), min(0.3,0.4), min(0.1,0.0)) = 0.8")
add_para("得出模糊结论 B' = [0.2, 0.5, 0.8]。")

add_para("【4. 第三步：解模糊判决（加权平均法）】", bold=True)
add_para("采用加权平均法得到确定数值：")
add_math_block("U = ∑(b'j * yj) / ∑b'j")
add_math_block("U = (0.2*20 + 0.5*50 + 0.8*80) / (0.2 + 0.5 + 0.8)")
add_math_block("U = (4 + 25 + 64) / 1.5 = 93 / 1.5 = 62%")
add_para("决策结果：模型建议今日买入标准仓位的 62%（约 300-400 元），而非全仓梭哈。")

# 2.2 遗传算法
doc.add_heading('2.2 应用二：基于遗传算法选择机制的资产配置', level=2)
add_para("在持有“证券、纳指、煤炭”三类资产时，我利用遗传算法中的轮盘赌选择逻辑筛选最优方案。")
add_para("【详细计算过程】", bold=True)
add_para("假设种群中有 3 个资产配置个体，其夏普比率（适应度 Fitness）如下：")
add_para("个体 A（侧重纳指）：f1 = 1.5；个体 B（侧重证券）：f2 = 0.8；个体 C（侧重煤炭）：f3 = 0.2")

add_para("1. 计算总适应度：Sum = 1.5 + 0.8 + 0.2 = 2.5")
add_para("2. 计算选择概率 (P = f / Sum)：")
add_para("   PA = 60%, PB = 32%, PC = 8%")
add_para("3. 计算累积概率 (q)：")
add_para("   qA = 0.60")
add_para("   qB = 0.60 + 0.32 = 0.92")
add_para("   qC = 0.92 + 0.08 = 1.00")
add_para("4. 随机选择（模拟轮盘赌）：")
add_para("   生成随机数 r = 0.75。")
add_para("   判断：因为 0.60 ≤ r < 0.92，系统选中 个体 B 进入下一代。")
add_para("   Python 实证：在我的 portfolio.ipynb 脚本中，经过 5000 次模拟，系统收敛出的最优解显示纳指与证券的配置比例较高，与此逻辑一致。")

# 2.3 感知机
doc.add_heading('2.3 应用三：基于感知机的行情趋势预测', level=2)
add_para("为了预测富国煤炭（013275）明日涨跌，我构建了一个单层感知机模型。")
add_para("【应用过程与权重迭代】", bold=True)
add_para("1. 特征定义：输入向量 x = [x1, x2]。x1为昨日涨跌幅(-1)，x2为乖离率(-2)。期望输出 y = 1 (涨)。")
add_para("2. 初始化：权重 ω = [0.5, 0.5]，阈值 θ = 0，学习率 η = 0.1。")
add_para("3. 前向计算（算 y 预测值）：")
add_math_block("S = 0.5*(-1) + 0.5*(-2) = -1.5")
add_math_block("f(S) = -1 (预测为跌)")
add_para("4. 误差修正（ω 迭代）：")
add_para("   实际情况：次日受寒潮影响超跌反弹，实际为涨 (y=1)。")
add_math_block("误差 e = 1 - (-1) = 2")
add_math_block("Δω = 0.1 * 2 * x")
add_math_block("ω1(new) = 0.5 + 0.2*(-1) = 0.3")
add_math_block("ω2(new) = 0.5 + 0.2*(-2) = 0.1")
add_para("结论：权重减小，模型学会了在“大乖离率”时降低惯性下跌的权重，从而预测反弹。")

# 5. 第三章 总结
doc.add_heading('三、 总结与体会', level=1)
add_para("本次报告将《人工智能导论》中的经典理论与我的数字经济专业研究进行了深度融合。")
add_para("1. 模糊推理通过矩阵运算，解决了金融交易中“买多少”的非线性决策问题。")
add_para("2. 遗传算法的轮盘赌机制，为资产组合寻找全局最优解提供了高效的路径。")
add_para("3. 感知机的权重迭代过程，直观展示了机器如何从“预测错误”中学习市场规律。")
add_para("未来，我计划进一步结合导师在生态经济与林业碳汇方向的研究，利用 Python 和更深层次的神经网络（如 LSTM），对福建省碳交易价格进行更精准的量化预测，探索“AI+绿色金融”的交叉应用。")

# 6. 保存文件
file_name = 'AI_Course_Report_FuQinghua.docx'
doc.save(file_name)
print(f"✅ 报告生成成功！文件名：{file_name}")